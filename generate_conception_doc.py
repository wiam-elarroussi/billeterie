import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("DOSSIER DE CONCEPTION TECHNIQUE — STACK, DATA & APIS")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Somayar S.A.R.L. AU — Technical Architecture Document — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_conception_doc():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Cover Section
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — ARCHITECTURE & INGENIERIE LOGICIELLE")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Amber
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("DOSSIER DE CONCEPTION TECHNIQUE : STACK, MODÉLISATION DES DONNÉES & DÉCOUPAGE DES APIS")
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("Spécifications architecturales de production pour le système global de billetterie, contrôle d'accès et cashless E-Ticket Pro (Normes FIFA Ready / Grand Stade de Casablanca)")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta Table
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta_data = [
        ("Projet / Marché", "Système de Billetterie & Accès — Grand Stade de Casablanca (AO N° 02/2026/GSC)"),
        ("Nature du Document", "Dossier de Conception Technique (DCT) & Spécifications des APIs"),
        ("Stack Cible", "Next.js 14, NestJS / Go, PostgreSQL 16, Redis Cluster, gRPC, Node Edge Dell R360"),
        ("Auteur & Intégrateur", "SOMAYAR S.A.R.L. AU — Direction Technique & R&D"),
        ("Date & Version", "Juillet 2026 — Version 1.0 Finale")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Formatting Helpers
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="ARCHITECTURAL DECISION RECORD (ADR)"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "EFF6FF") # Soft Blue
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(30, 64, 175)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_h1("1. CHOIX DE LA STACK TECHNIQUE & JUSTIFICATIONS ARCHITECTURALES")
    add_body("L'architecture technique de la solution E-Ticket Pro est conçue pour répondre aux contraintes d'infrastructures d'envergure internationale (type Grand Stade de Casablanca : 67 000 places, débit de 1000 spectateurs/minute par porte, normes FIFA Ready). L'architecture adoptée est une topologie Hybride Cloud / Edge Computing.")

    add_h2("1.1 Matrice Détaillée des Choix Technologiques & Justifications")

    # Table Stack Justification
    tbl_stack = doc.add_table(rows=8, cols=3)
    tbl_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_stack.autofit = False
    
    stack_data = [
        ("Brique Systémique", "Technologie Sélectionnée", "Justification Technique & Métier"),
        ("Frontend Web E-Commerce", "Next.js 14+ (React, TypeScript, TailwindCSS)", "Server-Side Rendering (SSR) pour un chargement instantané, rendu SEO parfait, intégration PWA native et support d'affichage réactif sur tous les mobiles."),
        ("Guichet POS (Caisses)", "Electron / PWA (TypeScript + WebSerial API)", "Communication directe avec les imprimantes thermiques de billets et lecteurs de cartes via l'API WebSerial, exécutable en mode autonome."),
        ("App Mobile PDA (Contrôle)", "React Native / Flutter + C++ Native Scanner", "Performance native sur terminaux portables durcis Android/iOS, gestion ultra-rapide de l'appareil photo/laser et stockage SQLite local déconnecté."),
        ("Backend & Microservices", "NestJS (TypeScript) & Go (Golang)", "Go offre une concurrence inégalée (Goroutines) pour absorber 10 000+ requêtes/sec lors des ouvertures de billetterie, NestJS apporte la rigueur architecturale."),
        ("Base de Données Principale", "PostgreSQL 16 + PostGIS Extension", "SGBDR robuste assurant le respect absolu des transactions ACID. L'extension PostGIS permet les requêtes géospatiales sur la disposition 2D des tribunes."),
        ("Cache & Verrouillage Sièges", "Redis Cluster (Memory Store)", "Verrouillage pessimiste ultra-rapide des sièges sélectionnés (< 2ms) pour empêcher deux acheteurs de réserver la même place simultanément."),
        ("Edge Computing Local (Stade)", "Dell PowerEdge R360 + Local Redis Engine", "Serveur local physique 1U au stade. Stocke le cache des billets pour assurer la validation en < 0.15s même en cas de coupure de la fibre Internet.")
    ]

    col_w = [Inches(1.8), Inches(2.0), Inches(2.7)]
    for idx, r_row in enumerate(stack_data):
        row = tbl_stack.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_w[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(9.0)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)

    add_callout("L'architecture sépare hermétiquement le traitement en ligne (Cloud Ventes) du traitement physique aux portes (Edge Local Stade). Même en cas d'attaque DDoS massive sur le site Web ou d'interruption du câble réseau principal, les tourniquets du stade continuent de valider les spectateurs à un rythme de 0,15s par billet grâce à la base locale répliquée sur le serveur Dell R360.", "RÉSILIENCE ARCHITECTURAL CLOUD / EDGE")

    # --- SECTION 2 ---
    add_h1("2. MODÉLISATION COMPLÈTE DES DONNÉES (SCHÉMA SQL & DICTIONNAIRE)")
    add_body("La base de données relationnelle (PostgreSQL 16) est modélisée selon les principes de la 3ème Forme Normale (3NF). Elle est structurée en 5 domaines fonctionnels majeurs.")

    # Table 1: Users & Auth
    add_h2("2.1 Domaine 1 : Gestion des Utilisateurs & Habilitations (IAM)")
    tbl_u = doc.add_table(rows=5, cols=4)
    tbl_u.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_u.autofit = False
    
    u_data = [
        ("Nom du Champ", "Type de Donnée", "Contraintes / Clés", "Description Fonctionnelle"),
        ("id", "UUID", "PRIMARY KEY, DEFAULT gen_random_uuid()", "Identifiant unique universel de l'utilisateur."),
        ("email", "VARCHAR(255)", "UNIQUE, NOT NULL", "Adresse email servant d'identifiant de connexion."),
        ("password_hash", "VARCHAR(255)", "NOT NULL", "Mot de passe chiffré via l'algorithme Argon2id / Bcrypt."),
        ("role", "ENUM", "NOT NULL, DEFAULT 'CLIENT'", "Rôle RBAC (SUPER_ADMIN, STADIUM_ADMIN, POS_AGENT, PDA_AGENT, CLIENT).")
    ]

    col_db = [Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.7)]
    for idx, r_row in enumerate(u_data):
        row = tbl_u.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_db[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(8.5)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=70, right=70)

    # Table 2: Venues & Seats
    add_h2("2.2 Domaine 2 : Infrastructure du Stade & Cartographie 2D (Seats)")
    tbl_s = doc.add_table(rows=6, cols=4)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_s.autofit = False
    
    s_data = [
        ("Nom du Champ", "Type de Donnée", "Contraintes / Clés", "Description Fonctionnelle"),
        ("id", "BIGINT", "PRIMARY KEY, GENERATED ALWAYS AS IDENTITY", "Identifiant séquentiel de l'emplacement."),
        ("venue_id", "UUID", "FOREIGN KEY -> Venues(id)", "Référence du stade (ex: Grand Stade de Casablanca)."),
        ("stand_name", "VARCHAR(50)", "NOT NULL", "Nom de la tribune (NORD, SUD, EST, OUEST, VIP)."),
        ("row_number", "INT", "NOT NULL", "Numéro de la rangée dans le secteur."),
        ("seat_number", "INT", "NOT NULL", "Numéro individuel du siège physique sur le rang.")
    ]

    for idx, r_row in enumerate(s_data):
        row = tbl_s.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_db[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(8.5)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=70, right=70)

    # Table 3: Tickets & Barcode
    add_h2("2.3 Domaine 3 : Billetterie, QR Codes & Checksum Mode 4 (Tickets)")
    tbl_t = doc.add_table(rows=7, cols=4)
    tbl_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_t.autofit = False
    
    t_data = [
        ("Nom du Champ", "Type de Donnée", "Contraintes / Clés", "Description Fonctionnelle"),
        ("id", "UUID", "PRIMARY KEY, DEFAULT gen_random_uuid()", "Identifiant unique du billet."),
        ("event_id", "UUID", "FOREIGN KEY -> Events(id)", "Référence de l'événement / match."),
        ("seat_id", "BIGINT", "FOREIGN KEY -> Seats(id)", "Référence du siège réservé."),
        ("buyer_id", "UUID", "FOREIGN KEY -> Users(id)", "Référence de l'acheteur enregistrer."),
        ("qr_hash", "VARCHAR(64)", "UNIQUE, NOT NULL", "Code unique Checksum Mode 4 infalsifiable (ex: ETK-2026-X9F4A7B2)."),
        ("status", "ENUM", "NOT NULL, DEFAULT 'ISSUED'", "Statut du billet (ISSUED, VALIDATED, CANCELLED, BLACKLISTED).")
    ]

    for idx, r_row in enumerate(t_data):
        row = tbl_t.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_db[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(8.5)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=70, right=70)

    # Table 4: Access Logs
    add_h2("2.4 Domaine 4 : Traçabilité des Scans d'Accès (AccessLogs)")
    tbl_l = doc.add_table(rows=6, cols=4)
    tbl_l.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_l.autofit = False
    
    l_data = [
        ("Nom du Champ", "Type de Donnée", "Contraintes / Clés", "Description Fonctionnelle"),
        ("id", "BIGINT", "PRIMARY KEY, GENERATED ALWAYS AS IDENTITY", "Identifiant du log de contrôle."),
        ("ticket_id", "UUID", "FOREIGN KEY -> Tickets(id)", "Billet présenté au scanner."),
        ("scanned_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()", "Horodatage au millième de seconde de la tentative."),
        ("gate_code", "VARCHAR(20)", "NOT NULL", "Code de la porte / tourniquet (ex: PORTE-NORD-02)."),
        ("result", "ENUM", "NOT NULL", "Résultat de la validation (GRANTED, REJECTED_DUPLICATE, REJECTED_INVALID).")
    ]

    for idx, r_row in enumerate(l_data):
        row = tbl_l.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_db[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(8.5)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=70, right=70)

    # --- SECTION 3 ---
    add_h1("3. SPÉCIFICATIONS ET DÉCOUPAGE EN APIS (RESTFUL & gRPC)")
    add_body("L'architecture logicielle découpe les fonctionnalités en 6 modules d'APIs spécialisées. Les échanges grand public et administration utilisent REST HTTP/2 JSON, tandis que les contrôles d'accès haute performance aux tourniquets utilisent gRPC Protocol Buffers.")

    # API 1
    add_h2("3.1 Service 1 : Auth & Identity Management API (`/api/v1/auth`)")
    add_bullet("Authentification et génération de jeton JWT d'accès + Refresh Token.", "POST /api/v1/auth/login : ")
    add_bullet("Inscription des nouveaux acheteurs grand public.", "POST /api/v1/auth/register : ")
    add_bullet("Lecture des informations et rôle de l'utilisateur connecté.", "GET /api/v1/auth/me : ")

    # API 2
    add_h2("3.2 Service 2 : Venue & Stadium 2D Layout API (`/api/v1/venues`)")
    add_bullet("Récupération de la structure 2D complète du stade (tribunes, secteurs).", "GET /api/v1/venues/{id}/layout : ")
    add_bullet("Retourne l'état d'occupation en temps réel de tous les sièges d'une tribune pour un événement.", "GET /api/v1/venues/{id}/tribunes/{name}/seats?event_id={id} : ")
    add_bullet("Verrouillage temporaire d'un siège pendant 5 minutes (Redis Lock) pendant que le client effectue le paiement.", "POST /api/v1/venues/seats/lock : ")

    # API 3
    add_h2("3.3 Service 3 : Event & Ticketing Engine API (`/api/v1/events` & `/tickets`)")
    add_bullet("Liste des événements à venir avec filtres (date, statut, tarifs).", "GET /api/v1/events : ")
    add_bullet("Création d'un événement et configuration des grilles tarifaires.", "POST /api/v1/events (Admin) : ")
    add_bullet("Achat et émission de billets avec génération du QR Code Checksum Mode 4.", "POST /api/v1/tickets/purchase : ")
    add_bullet("Récupération de la liste des billets achetés par le client connecté (Wallet).", "GET /api/v1/tickets/my-tickets : ")

    # API 4
    add_h2("3.4 Service 4 : Point of Sale POS API (`/api/v1/pos`)")
    add_bullet("API hautement optimisée pour la vente physique en 3 clics sur caisse tactile.", "POST /api/v1/pos/express-checkout : ")
    add_bullet("Clôture de caisse et émission du rapport des encaissements par caissier.", "GET /api/v1/pos/shift-report : ")

    # API 5
    add_h2("3.5 Service 5 : Access Control gRPC Service (`AccessControlService`)")
    add_body("Service binaire haute performance gRPC exécuté en local sur le serveur Dell R360 du stade pour asservir les tourniquets et PDA agents :")
    add_bullet("RPC de validation d'un QR Code en < 0.15s avec retour du résultat (GRANTED, REJECTED_DUPLICATE, REJECTED_INVALID).", "rpc ValidateTicket (ValidateRequest) returns (ValidateResponse) : ")
    add_bullet("Flux bidirectionnel continu (gRPC Streaming) pour resynchroniser les registres locaux d'accès avec le serveur cloud dès que la liaison réseau est rétablie.", "rpc SyncOfflineLogs (stream LogChunk) returns (SyncResult) : ")

    # API 6
    add_h2("3.6 Service 6 : Smart Reporting & Analytics API (`/api/v1/analytics`)")
    add_bullet("Synthèse temps réel des jauges de remplissage par porte et du chiffre d'affaires.", "GET /api/v1/analytics/live-summary : ")
    add_bullet("Génération et téléchargement des rapports financiers et d'accès aux formats CSV, PDF et Excel.", "GET /api/v1/analytics/export?format=csv : ")

    # --- SECTION 4 ---
    add_h1("4. STRATÉGIE DE SYNCHRONISATION LOCAL / CLOUD (OFFLINE ENGINE)")
    add_body("Pour garantir la conformité aux exigences FIFA Ready et assurer qu'aucun tourniquet ne soit bloqué en cas de panne réseau, le système intègre un moteur de synchronisation déconnecté :")
    add_bullet("Avant le début de chaque événement (H-4), le serveur local Dell R360 du stade télécharge l'intégralité des identifiants et Checksums des billets émis dans sa base Redis/RocksDB locale.", "Pré-Chargement des Billets : ")
    add_bullet("Pendant l'événement, les validations sont effectuées à 100% sur la base locale du stade avec un temps de réponse < 0,15s.", "Validation Locale Autonome : ")
    add_bullet("Les logs de passage sont accumulés dans une file d'attente locale sérialisée.", "File d'Attente Locale : ")
    add_bullet("Dès rétablissement du lien VPN IPSEC, un agent CDC (Change Data Capture) déverse en tâche de fond les logs accumulés vers la base PostgreSQL cloud sans impacter les tourniquets.", "Resynchronisation Transparente : ")

    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Dossier_Conception_Technique_Stack_Data_API.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_conception_doc()
