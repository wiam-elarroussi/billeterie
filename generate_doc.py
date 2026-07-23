import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("CAHIER DES CHARGES — SOLUTION E-TICKET PRO | CONFIDENTIEL")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Somayar S.A.R.L. AU — Project FIFA Ready / Grand Stade de Casablanca — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_cahier_des_charges():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Title Page / Cover Section
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — EQUIPEMENTS DE CONTRÔLE & DE SÉCURITÉ")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Accent Gold/Orange
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("CAHIER DES CHARGES TECHNIQUE & FONCTIONNEL")
    r_title.font.bold = True
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(30)
    r_sub = p_subtitle.add_run("Solution globale et intégrée de Billetterie Informatisée, Paiement Cashless, Contrôle d'Accès Haute Performance et Smart Reporting (« E-Ticket Pro / FIFA Ready »)")
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta Box
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta_data = [
        ("Projet / Référence", "Appel d'Offres N° 02/2026/GSC — Grand Stade de Casablanca"),
        ("Solution Système", "Écosystème E-Ticket Pro (E-Ticket Pay, E-Ticket Report, Central Controller)"),
        ("Conformité", "Standards FIFA Ready, ISO 27001, Mode Autonome / Offline (0,5s/validation)"),
        ("Auteur / Éditeur", "SOMAYAR S.A.R.L. AU — Zone Industrielle Ain Atiq, Témara"),
        ("Date & Version", "Juillet 2026 — Version 1.0 Officielle")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Helper functions for sections
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="NOTE REQUISITION FIFA READY"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "FEF3C7") # Soft amber
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(180, 83, 9)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_heading_1("1. PRÉSENTATION DU PROJET & CONTEXTE OPERATIONNEL")
    add_body("Le présent Cahier des Charges définit les exigences techniques, fonctionnelles, architecturales et opérationnelles pour le déploiement d'une solution de billetterie informatisée tout-en-un, de gestion des accès et de paiement cashless dématérialisé. Ce système est spécifiquement dimensionné pour répondre aux exigences des grands complexes sportifs et événementiels de niveau international, notamment dans le cadre des appels d'offres internationaux type Grand Stade de Casablanca (Appel d'Offres N° 02/2026/GSC).")
    
    add_heading_2("1.1 Objectifs Stratégiques")
    add_bullet("Assurer la gestion centralisée et sécurisée de la billetterie multi-site et multi-événement en temps réel (vente en ligne, guichets fixes, terminaux mobiles).", "Billetterie Tout-en-un : ")
    add_bullet("Offrir une expérience fluide et sans attente aux accès avec un temps de validation de chaque titre d'accès inférieur à 0,5 seconde.", "Contrôle d'Accès Ultra-Rapide : ")
    add_bullet("Intégrer une solution de paiement dématérialisé Cashless et d'appairage multi-supports (NFC, QR Code, Barcode 1D/2D, Apple Pay, Google Pay).", "Paiement Cashless & E-Commerce : ")
    add_bullet("Fournir des outils décisionnels et des tableaux de bord dynamiques (Smart-Reporting) accessibles 24/7 sur PC et appareils mobiles (Android/iOS).", "Supervision & Smart Reporting : ")
    add_bullet("Garantir une conformité absolue aux exigences technologiques et opérationnelles des compétitions internationales FIFA Ready.", "Conformité FIFA Ready : ")

    add_heading_2("1.2 Périmètre Global de la Solution")
    add_body("L'écosystème logiciel et matériel comprend 5 modules principaux interconnectés :")
    add_bullet("Module cœur d'administration générale, création des stades, agencement 2D des tribunes/sièges, gestion des événements et des tarifs.", "E-Ticket Pro (Core System) : ")
    add_bullet("Plateforme de vente en ligne (e-boutique intégrée / API Cloud) et module guichets tactiles POS pour la vente physique sur site.", "E-Ticket Pay (Cashless & Sales) : ")
    add_bullet("Moteur d'analyse décisionnelle, statistiques des ventes, comportement d'achat et génération de rapports multi-formats (PDF, Excel, Word, CSV, XML).", "E-Ticket Report (Smart Analytics) : ")
    add_bullet("Matériel serveur rackable haute disponibilité (Dell PowerEdge R360) hébergeant les bases de données et les moteurs de validation.", "Serveur Contrôleur Central : ")
    add_bullet("Attestation de conformité, API & SDK FIFA Ticketing / Ticketmaster, gestion des accréditations (VIP, Presse, Délégations) et mode déconnecté résilient.", "Module d'Interopérabilité FIFA : ")

    # --- SECTION 2 ---
    add_heading_1("2. ARCHITECTURE TECHNIQUE & INFRASTRUCTURE RÈSEAU/MATÉRIEL")
    add_body("L'architecture du système E-Ticket Pro repose sur une topologie hybride (Serveurs locaux de haute sécurité couplés à des serveurs Cloud synchronisés via tunnel VPN IPSEC chiffré), garantissant zéro perte de données et une résilience totale en cas de coupure réseau.")

    add_heading_2("2.1 Spécifications du Serveur Central de Contrôle d'Accès et Billetterie")
    add_body("Le système s'appuie sur une infrastructure serveur Dell PowerEdge R360 (Rack 1U) certifiée entreprise, dont les caractéristiques techniques minimales requises sont récapitulées ci-dessous :")

    # Table for Dell R360
    tbl_server = doc.add_table(rows=11, cols=2)
    tbl_server.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_server.autofit = False
    
    server_specs = [
        ("Composant / Caractéristique", "Spécification Technique Requise (Dell PowerEdge R360)"),
        ("Processeur (CPU)", "Intel® Xeon® E-2400 ou Xeon® 6300 (jusqu'à 8 cœurs) / Pentium G7400"),
        ("Mémoire Vive (RAM)", "4 emplacements DDR5 DIMM, jusqu'à 128 GB max UDIMM ECC 4400 MT/s"),
        ("Contrôleur de Stockage", "PERC H755 / H355 Adpt, Boot optimisé BOSS-N1 (HWRAID 2x M.2 SSD NVMe)"),
        ("Baies de Disques", "Jusqu'à 4 x 3.5\" SAS/SATA (max 64 TB) ou 8 x 2.5\" SAS/SATA SSD (max 61.44 TB)"),
        ("Alimentation Réduite", "Double alimentation redondante 600W Platinum / 700W Titanium Hot-Swap (100-240V)"),
        ("Gestion Embarquée", "iDRAC9 Enterprise, iDRAC Direct, RESTful API Redfish, Service Manual"),
        ("Réseau & Connectivity", "2 x 1 GbE LOM intégrés, 1 x Port Ethernet iDRAC dédié, USB 3.2 Gen1, VGA, Serial"),
        ("Sécurité Matérielle", "Silicon Root of Trust, Secure Boot, Chiffrement des données à l'arrêt (SEDs), TPM 2.0 FIPS"),
        ("Systèmes d'Exploitation", "Canonical Ubuntu Server LTS, Red Hat Enterprise Linux, VMware ESXi, Windows Server"),
        ("Facteur de Forme", "Serveur Rack 1U haute densité pour baie informatique sécurisée")
    ]

    for idx, (c1_txt, c2_txt) in enumerate(server_specs):
        row = tbl_server.rows[idx]
        cell_a, cell_b = row.cells[0], row.cells[1]
        cell_a.width = Inches(2.2)
        cell_b.width = Inches(4.3)
        
        pa = cell_a.paragraphs[0]
        pb = cell_b.paragraphs[0]
        
        ra = pa.add_run(c1_txt)
        rb = pb.add_run(c2_txt)
        
        if idx == 0:
            ra.font.bold = True
            ra.font.color.rgb = RGBColor(255, 255, 255)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell_a, "1B365D")
            set_cell_background(cell_b, "1B365D")
        else:
            ra.font.bold = True
            ra.font.size = Pt(9.5)
            ra.font.color.rgb = RGBColor(27, 54, 93)
            rb.font.size = Pt(9.5)
            bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell_a, bg)
            set_cell_background(cell_b, bg)
            
        set_cell_margins(cell_a, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cell_b, top=80, bottom=80, left=100, right=100)

    add_heading_2("2.2 Infrastructure Réseau et Équipements de Terrain")
    add_bullet("Connexion chiffrée permanente via VPN IPSEC entre le serveur Cloud de vente en ligne et le contrôleur local du stade.", "Tunnel VPN IPSEC Sécurisé : ")
    add_bullet("Commutateurs réseau fédérateurs et locaux administrables assurant la liaison Très Haut Débit entre le PC de sécurité, les guichets POS et les bornes d'accès.", "Switchs Fédérateurs & Access Points Wi-Fi : ")
    add_bullet("Tourniquets de sécurité, portillons pivotants et couloirs rapides équipés d'unités d'asservissement autonomes et de lecteurs multitechnologies.", "Obstacles Physiques d'Accès : ")
    add_bullet("Terminaux portables durcis (PDA) communicant en Wi-Fi/4G/5G, dotés de scanners laser 1D/2D et puces NFC pour les agents de contrôle itinérants.", "Terminaux Mobiles PDA : ")
    add_bullet("Postes de caisse tactiles (PC + Imprimante thermique + Douchette/Lecteur + Connexion Internet/LAN) assurant une vente billetterie en 3 clics.", "Points de Vente Locaux (POS) : ")

    add_callout("Le système intègre un mode déconnecté (Offline) avancé. En cas d'interruption du réseau physique ou Wi-Fi, les unités de traitement local des obstacles d'accès et les PDA conservent l'intégralité des droits d'accès en mémoire cache local et valident les billets en autonomie absolue. Dès le rétablissement du réseau, les données de comptage et de validation sont resynchronisées automatiquement avec le serveur central sans moindre perte de données.", "EXIGENCE CRITIQUE : CONTINUITÉ DE SERVICE EN MODE OFFLINE")

    # --- SECTION 3 ---
    add_heading_1("3. SPÉCIFICATIONS FONCTIONNELLES DÉTAILLÉES")
    
    add_heading_2("3.1 Module E-Ticket Pro — Gestion de la Billetterie et Configuration")
    add_body("Le module E-Ticket Pro est l'outil central de paramétrage de l'écosystème. Ses fonctionnalités couvrent :")
    add_bullet("Création fonctionnelle et représentation graphique interactive 2D des stades et enceintes (bâtiments, portes, accès, tribunes, zones, secteurs, rangs et sièges).", "Éditeur de Plan d'Enceinte 2D : ")
    add_bullet("Possibilité de numéroter les sièges avec définition automatique du sens de passage de siège à siège.", "Numérotation & Sens de Passage : ")
    add_bullet("Gestion complète des compétitions, événements, équipes, catégories de billets, grille tarifaire dynamique (par place, catégorie de visiteurs, saison, VIP, etc.).", "Gestion Événementielle & Tarification : ")
    add_bullet("Création de groupes permanents pour les abonnés (saisonniers, éliminatoires, packs) avec activation spécifique de l'accès aux événements programmés.", "Gestion des Abonnements & Packs : ")
    add_bullet("Attribution des vendeurs externes à des canaux de vente spécifiques, définition des règles de vente, des horaires d'ouverture/fermeture des billetteries.", "Canaux de Vente & Quotas : ")
    add_bullet("Définition dynamique des quotas d'impression par client (ex: maximum 3 billets par acheteur) et quotas affectés à chaque point de vente en temps réel.", "Contrôle Strict des Quotas : ")
    add_bullet("Studio de création graphique de la charte visuelle des billets thermiques, e-tickets PDF et badges d'accréditation.", "Conception Visuelle des Billets : ")
    add_bullet("Modules de réservation de groupe, gestion des conventions d'entreprises, réimpression de commandes et gestion des annulations.", "Conventions & Annulations : ")

    add_heading_2("3.2 Module E-Ticket Pay — Vente, Cashless & E-Commerce")
    add_body("Le module E-Ticket Pay offre une solution intégrée pour dématérialiser les transactions et fluidifier les ventes :")
    add_bullet("Intégration transparente sur le site web officiel en tant que boutique en ligne ou site de billetterie indépendant.", "Vente en Ligne E-Commerce : ")
    add_bullet("Interface ultra-ergonomique adaptée aux écrans tactiles des guichets, permettant d'effectuer la vente physique en 3 clics seulement.", "Guichet POS Tactile Express : ")
    add_bullet("Système de paiement sans contact (Cashless) permettant le rechargement de compte en ligne ou aux guichets, l'appairage direct du solde avec le billet/badge dès l'entrée.", "Module Cashless & Rechargement : ")
    add_bullet("Permet aux spectateurs d'enregistrer et de visualiser la géolocalisation exacte de leur siège sur une carte interactive 2D sur le navigateur de leur mobile.", "Enregistrement & Repérage des Sièges : ")
    add_bullet("Gestion des bons d'achat, des coupons de réduction, des prestations VIP et du catering précommandé.", "Vouchers & Précommandes : ")
    add_bullet("Création d'un espace personnel pour les participants, collecte des préférences d'achat et envoi de notifications/courriels ciblés.", "Espace Client & Marketing : ")

    add_heading_2("3.3 Module E-Ticket Report — Smart Reporting & Analytics")
    add_body("Le module E-Ticket Report fournit un pilotage décisionnel en temps réel 24h/24 et 7j/7 :")
    add_bullet("Tableaux de bord graphiques affichant en temps réel la jauge d'occupation globale, le flux de passage par porte et les statistiques de vente par catégorie.", "Tableaux de Bord Dynamiques : ")
    add_bullet("Visualisation des revenus de billetterie, comparaisons instantanées avec des événements ou périodes antérieures.", "Analyse Financière & Revenus : ")
    add_bullet("Rapports d'acquisition (URL de référence, canaux marketing), répartition géographique régionale des acheteurs, données démographiques (âge, genre, intérêts).", "Analyse Démographique & Client : ")
    add_bullet("Sauvegarde automatique et archivage permanent de l'historique de toutes les éditions, réservations et transactions financières.", "Archivage & Historique : ")
    add_bullet("Exportation en 1 clic de tous les rapports et listes au format Excel, Word, CSV, PDF, XML, ainsi que sous forme de graphiques visuels.", "Multi-Formats d'Exportation : ")
    add_bullet("Application mobile de supervision en temps réel compatible avec les appareils Android, iOS et macOS pour les responsables de sécurité et organisateurs.", "Interface Supervision Mobile : ")

    add_heading_2("3.4 Spécifications d'Interopérabilité et Normes FIFA Ready")
    add_body("Pour répondre aux exigences des compétitions internationales organisées sous l'égide de la FIFA (notamment pour le Grand Stade de Casablanca), la solution E-Ticket Pro garantit :")
    add_bullet("Fourniture complète de jeux d'API RESTful et SDK sécurisés pour un couplage natif avec les systèmes de billetterie FIFA Ticketing ou Ticketmaster.", "API & SDK d'Intégration FIFA : ")
    add_bullet("Traitement et validation d'un titre d'accès en moins de 0,5 seconde aux obstacles d'entrée pour maintenir un débit optimal aux tourniquets.", "Vitesse de Validation < 0,5s : ")
    add_bullet("Prise en charge native des codes-barres 1D/2D (imprimés ou sur smartphone), des puces sans contact NFC, ainsi que des portefeuilles virtuels Apple Pay et Google Pay.", "Multi-Support d'Accès : ")
    add_bullet("Gestion centralisée et transparente des accréditations pour les populations spécifiques (VIP, VVIP, Presse, Bénévoles, Délégations officielles, Staff).", "Gestion des Accréditations : ")
    add_bullet("Génération de codes uniques sécurisés infalsifiables intégrant des caractères de Checksum en Mode 4, avec filtrage par liste noire et liste blanche.", "Sécurité & Anti-Fraude (Checksum Mode 4) : ")
    add_bullet("Supervision et contrôle centralisé de tous les obstacles physiques (tourniquets, couloirs rapides) directement depuis le poste de sécurité central (PC Sécurité).", "Monitoring Centralisé des Obstacles : ")

    # --- SECTION 4 ---
    add_heading_1("4. SÉCURITÉ ET MATRICE DES RÔLES & PRIVILÈGES")
    add_body("Le système E-Ticket Pro est conçu sur un modèle d'habilitation basé sur les rôles (RBAC - Role-Based Access Control). Chaque utilisateur créé dans le système se voit attribuer un jeu précis de privilèges.")

    # Table of Roles
    tbl_roles = doc.add_table(rows=7, cols=2)
    tbl_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_roles.autofit = False
    
    roles_data = [
        ("Profil Utilisateur / Rôle", "Périmètre des Droits et Privilèges Accordés"),
        ("Super Administrateur", "Accès total au système, gestion de la matrice des rôles, configuration serveur, audits de sécurité et logs globaux."),
        ("Administrateur Stade", "Création des enceintes, plans 2D, gestion des compétitions, événements, grille tarifaire et affectation des quotas."),
        ("Superviseur PC Sécurité", "Supervision en temps réel des jauges, monitoring graphique des obstacles d'accès, gestion des listes noires/blanches et alertes d'intrusion."),
        ("Opérateur Guichet / Caissier", "Vente rapide de billets (3 clics), encaissement Cashless/CB/Espèces, réimpression autorisée de commandes, édition de factures."),
        ("Agent de Contrôle (PDA)", "Validation des titres d'accès aux portes via scanner/NFC, contrôle des accréditations, basculement automatique en mode déconnecté."),
        ("Client / Participant", "Consultation des événements, achat en ligne, rechargement du compte Cashless, consultation du repérage 2D du siège, historique des billets.")
    ]

    for idx, (r_title, r_desc) in enumerate(roles_data):
        row = tbl_roles.rows[idx]
        ca, cb = row.cells[0], row.cells[1]
        ca.width = Inches(2.0)
        cb.width = Inches(4.5)
        
        pa, pb = ca.paragraphs[0], cb.paragraphs[0]
        ra, rb = pa.add_run(r_title), pb.add_run(r_desc)
        
        if idx == 0:
            ra.font.bold = True
            ra.font.color.rgb = RGBColor(255, 255, 255)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(ca, "1B365D")
            set_cell_background(cb, "1B365D")
        else:
            ra.font.bold = True
            ra.font.size = Pt(9.5)
            ra.font.color.rgb = RGBColor(27, 54, 93)
            rb.font.size = Pt(9.5)
            bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
            set_cell_background(ca, bg)
            set_cell_background(cb, bg)
            
        set_cell_margins(ca, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cb, top=80, bottom=80, left=100, right=100)

    # --- SECTION 5 ---
    add_heading_1("5. EXIGENCES NON-FONCTIONNELLES & QUALITÉ DE SERVICE (SLA)")
    add_bullet("Taux de disponibilité globale du système central supérieur ou égal à 99,99% pendant les événements.", "Haute Disponibilité (99.99%) : ")
    add_bullet("Temps de validation d'un billet aux obstacles physiques <= 0.5 sec ; temps d'affichage des tableaux de bord < 1.0 sec.", "Performance & Temps de Réponse : ")
    add_bullet("Sauvegarde continue des bases de données de billetterie avec réplication synchrone sur serveur de backup local et cloud.", "Sauvegarde & Intégrité : ")
    add_bullet("Conformité aux recommandations RGPD (protection des données personnelles des acheteurs) et ISO 27001.", "Protection des Données (RGPD) : ")

    # --- SECTION 6 ---
    add_heading_1("6. LIVRABLES DU PROJET & ENGAGEMENTS DE LIVRAISON")
    add_body("Dans le cadre du marché, SOMAYAR S.A.R.L. AU s'engage à fournir les livrables suivants :")
    add_bullet("Serveur rackable Dell PowerEdge R360 configuré, tourniquets/obstacles physiques, terminaux mobiles PDA, caisses tactiles POS et imprimantes thermiques.", "Infrastructure Matérielle : ")
    add_bullet("Licences d'utilisation des modules E-Ticket Pro, E-Ticket Pay, E-Ticket Report, APIs/SDKs d'interfaçage FIFA Ticketing.", "Suite Logicielle & Licences : ")
    add_bullet("Manuel d'exploitation réseau et système, guide d'utilisation des interfaces d'administration, guichets et PDA, documentation technique des APIs.", "Documentation Technique & Fonctionnelle : ")
    add_bullet("Attestation officielle de conformité aux exigences FIFA Ready et rapport d'essais en charge de validation à 0,5s/billet.", "Attestation de Conformité FIFA : ")
    add_bullet("Sessions de formation pour les administrateurs système, les superviseurs du PC Sécurité, les caissiers et les agents de contrôle aux accès.", "Plan de Formation & Transfert de Compétences : ")

    # Signature Block
    add_heading_2("Validation et Approbation du Cahier des Charges")
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_sig.autofit = False
    
    cell_s0 = tbl_sig.rows[0].cells[0]
    cell_s1 = tbl_sig.rows[0].cells[1]
    cell_s0.width = Inches(3.25)
    cell_s1.width = Inches(3.25)
    
    p_s0 = cell_s0.paragraphs[0]
    p_s0.add_run("Pour le Maître d'Ouvrage / Client :\n\n___________________________________\nDate & Signature").font.size = Pt(9.5)
    
    p_s1 = cell_s1.paragraphs[0]
    p_s1.add_run("Pour le Prestataire (SOMAYAR S.A.R.L. AU) :\n\n___________________________________\nDirection Technique E-Ticket Pro").font.size = Pt(9.5)
    
    set_cell_background(cell_s0, "F8FAFC")
    set_cell_background(cell_s1, "F8FAFC")
    set_cell_margins(cell_s0, top=120, bottom=120, left=120, right=120)
    set_cell_margins(cell_s1, top=120, bottom=120, left=120, right=120)

    # Header and Footer
    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Cahier_des_Charges_E-Ticket_Pro.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_cahier_des_charges()
