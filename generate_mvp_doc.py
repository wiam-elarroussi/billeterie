import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("SPÉCIFICATIONS FONCTIONNELLES MVP — E-TICKET PRO EXPRESS")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Somayar S.A.R.L. AU — Solution E-Ticket Pro / MVP Implementation Guide — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_mvp_features_doc():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Title Page / Cover Section
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — INGÉNIERIE & DÉVELOPPEMENT LOGICIEL")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Amber/Gold
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("SPÉCIFICATIONS DES FONCTIONNALITÉS DU MVP (MINIMUM VIABLE PRODUCT)")
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.paragraph_format.space_after = Pt(30)
    r_sub = p_subtitle.add_run("Guide opérationnel des fonctionnalités à concrétiser pour le MVP de la solution E-Ticket Pro (Billetterie, Guichet POS, Contrôle d'Accès Scanner & Smart Dashboard)")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta Box
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta_data = [
        ("Nom du Projet MVP", "E-Ticket Pro Express (MVP Ready)"),
        ("Objectif Technique", "Concrétisation rapide d'une version de démonstration et de validation fonctionnelle"),
        ("Périmètre Fonctionnel", "Auth RBAC, Stade 2D, Billetterie QR Code, Caisse POS 3-clics, Scanner PDA, Dashboard"),
        ("Éditeur / Développeur", "SOMAYAR S.A.R.L. AU"),
        ("Date & Version", "Juillet 2026 — Version 1.0 MVP Specification")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Helper formatting functions
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="DIRECTIVE TECH & OBJECTIF MVP"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "F0FDF4") # Soft green
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(22, 101, 52)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_h1("1. PÉRIMÈTRE ET STRATÉGIE DU MVP")
    add_body("L'objectif de ce document est de définir l'ensemble des fonctionnalités à concrétiser prioritairement dans le cadre du MVP (Minimum Viable Product) de l'application E-Ticket Pro. Le MVP permettra de valider techniquement et de démontrer l'efficacité des modules critiques : la gestion des événements, la réservation 2D, la vente physique en caisse express, la validation scanner par QR Code et le suivi en temps réel de la jauge d'occupation.")

    add_h2("1.1 Périmètre In-Scope vs Out-of-Scope")

    # Table In Scope vs Out of Scope
    tbl_scope = doc.add_table(rows=6, cols=2)
    tbl_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_scope.autofit = False
    
    scope_data = [
        ("Fonctionnalités In-Scope (Incluses dans le MVP)", "Fonctionnalités Out-of-Scope (V2 / Évolutions futures)"),
        ("Authentification & Matrice de 4 Rôles (Admin, Caissier, Agent PDA, Spectateur).", "Gestion poussée du matériel physique tourniquets/portillons (asservissement matériel)."),
        ("Visualiseur 2D dynamique de stade avec choix interactif des sièges.", "Module Cashless physique NFC avec terminal de rechargement sur carte bancaire."),
        ("Création d'événements, grilles tarifaires et quotas de vente.", "Intégration d'APIs tierces complexes (Ticketmaster, FIFA Ticketing SDK)."),
        ("Guichet POS tactile express pour la vente physique en 3 clics.", "Module marketing avancé (campagnes d'e-mailing ciblées, relances d'abandon de panier)."),
        ("Module Scanner de validation par caméra/PDA avec gestion du mode Offline.", "Gestion multi-stades complexe et réplication de bases distribuées multi-régions.")
    ]

    for idx, (in_txt, out_txt) in enumerate(scope_data):
        row = tbl_scope.rows[idx]
        ca, cb = row.cells[0], row.cells[1]
        ca.width = Inches(3.25)
        cb.width = Inches(3.25)
        pa, pb = ca.paragraphs[0], cb.paragraphs[0]
        ra, rb = pa.add_run(in_txt), pb.add_run(out_txt)
        if idx == 0:
            ra.font.bold = True
            ra.font.color.rgb = RGBColor(255, 255, 255)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(ca, "1B365D")
            set_cell_background(cb, "1B365D")
        else:
            ra.font.size = Pt(9.5)
            rb.font.size = Pt(9.5)
            bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
            set_cell_background(ca, bg)
            set_cell_background(cb, bg)
        set_cell_margins(ca, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cb, top=80, bottom=80, left=100, right=100)

    # --- SECTION 2 ---
    add_h1("2. DESCRIPTION DÉTAILLÉE DES MODULES DU MVP")

    # Module 1
    add_h2("2.1 Module 1 : Authentification & Gestion des Rôles (Auth & RBAC)")
    add_bullet("Formulaire de connexion épuré par identifiant/mot de passe avec gestion de jeton de session sécurisé.", "F01.1 - Connexion Sécurisée : ")
    add_bullet("Redirection automatique vers l'interface dédiée selon le profil connecté :", "F01.2 - Matrice de Redirection par Rôle : ")
    add_bullet("Accès au tableau de bord général, création d'événements, configuration du stade et gestion des utilisateurs.", "• Admin Stade : ")
    add_bullet("Accès direct à l'interface tactile POS de vente rapide de billets.", "• Agent Caissier : ")
    add_bullet("Accès à la caméra/scanner de validation des billets aux portes.", "• Agent PDA / Contrôleur : ")
    add_bullet("Accès au catalogue des matchs, au panier et à l'espace e-ticket.", "• Spectateur : ")

    # Module 2
    add_h2("2.2 Module 2 : Visualiseur & Plan 2D Interactif de Stade")
    add_bullet("Affichage schématique et esthétique 2D de l'enceinte (Tribunes Nord, Sud, Est, Ouest / VIP).", "F02.1 - Vue Générale d'Enceinte : ")
    add_bullet("Zoom sur une tribune spécifique avec rendu des blocs de sièges sous forme de grille interactive.", "F02.2 - Plan de Tribune & Sièges : ")
    add_bullet("Couleur verte = Libres, Couleur Rouge = Occupés/Vendus, Couleur Bleue = Sélectionnés.", "F02.3 - Code Couleur des Sièges : ")
    add_bullet("Au survol ou au clic sur un siège : affichage du numéro de rang, numéro de siège, tribune et tarif associé.", "F02.4 - Infobulle & Détails du Siège : ")

    # Module 3
    add_h2("2.3 Module 3 : Gestion des Événements & Générateur de Billets QR Code")
    add_bullet("Formulaire de création d'un événement (Nom du match/concert, Date, Heure, Description, Stade).", "F03.1 - Création d'Événement : ")
    add_bullet("Définition du prix unitaire par catégorie de tribune (ex: VIP = 500 DH, Tribune Est = 100 DH).", "F03.2 - Grille Tarifaire MVP : ")
    add_bullet("Pour chaque billet réservé, le système génère un identifiant unique incrackable et un QR Code au format SVG/PNG avec Checksum Mode 4.", "F03.3 - Moteur QR Code Checksum Mode 4 : ")
    add_bullet("Génération dynamique du billet imprimable/téléchargeable au format PDF synthétique avec QR Code, visuel et informations du match.", "F03.4 - Générateur PDF & E-Ticket : ")

    # Module 4
    add_h2("2.4 Module 4 : Interface de Caisse POS Tactile Express (3 Clics)")
    add_bullet("Interface optimisée pour écran tactile permettant de finaliser une vente guichet en 3 clics :", "F04.1 - Processus Vente Express : ")
    add_bullet("Clic 1 : Sélection de l'événement et de la tribune.", "• Clic 1 : ")
    add_bullet("Clic 2 : Choix de la quantité de billets ou sélection sur plan 2D.", "• Clic 2 : ")
    add_bullet("Clic 3 : Encaissement (Espèces/CB) et validation immédiate.", "• Clic 3 : ")
    add_bullet("Envoi direct du billet vers l'imprimante thermique ou affichage à l'écran du QR Code pour scan rapide par le client.", "F04.2 - Impression / Émission Billet : ")

    # Module 5
    add_h2("2.5 Module 5 : Boutique Web Client & Expérience Mobile Spectateur")
    add_bullet("Interface web responsive affichant la liste des matchs à venir avec compte à rebours et tarifs.", "F05.1 - Catalogue d'Événements : ")
    add_bullet("Sélection dynamique des places sur le plan 2D interactif depuis un navigateur smartphone.", "F05.2 - Réservation 2D Mobile : ")
    add_bullet("Module de simulation de paiement en ligne (CB) et confirmation instantanée.", "F05.3 - Checkout & Confirmation : ")
    add_bullet("Section 'Mes Billets' affichant la carte 2D de géolocalisation du siège et le QR Code à présenter au tourniquet.", "F05.4 - E-Ticket Mobile & Repérage Siège : ")

    # Module 6
    add_h2("2.6 Module 6 : Scanner de Validation aux Accès (Web PDA & Caméra)")
    add_bullet("Interface web/mobile exploitant la caméra du smartphone/PDA ou un scanner physique pour lire les QR Codes.", "F06.1 - Lecture QR Code : ")
    add_bullet("Temps de réponse de la validation < 0,5 seconde avec retour visuel et sonore immédiat :", "F06.2 - Rendu du Résultat (< 0.5s) : ")
    add_bullet("Vert + Bip aigu = Billet valide, accès autorisé.", "• ACCÈS ACCORDÉ : ")
    add_bullet("Rouge + Bip grave = Billet déjà scanné (Doublon) ou invalide.", "• ACCÈS REFUSÉ : ")
    add_bullet("Stockage local (LocalStorage / IndexedDB) des QR Codes valides de l'événement permettant de poursuivre les scans même sans connexion internet.", "F06.3 - Autonomie Offline (Mode Déconnecté) : ")

    # Module 7
    add_h2("2.7 Module 7 : Tableau de Bord & Monitoring (PC Sécurité / Supervision)")
    add_bullet("Jauge circulaire de taux de remplissage du stade mise à jour en temps réel.", "F07.1 - Jauge Remplissage Temps Réel : ")
    add_bullet("Compteurs dynamiques : Billets Vendus, Spectateurs Entrés (Scannés), Places Restantes.", "F07.2 - Compteurs de Flux : ")
    add_bullet("Histogramme des flux d'entrées par tranche de 15 minutes.", "F07.3 - Graphique des Entrées : ")
    add_bullet("Exportation en 1 clic de la liste des accès au format CSV et du bilan global en PDF.", "F07.4 - Export de Rapport : ")

    add_callout("L'application MVP sera développée sous la forme d'une Web App moderne (Single Page Application responsive) utilisant HTML5/Javascript (React/Vite) et CSS moderne (TailwindCSS/Vanilla CSS), offrant une interface fluide, animée et ultra-réactive pour les tests utilisateurs.", "STACK TECHNIQUE RECOMMANDÉE POUR LE DÉVELOPPEMENT")

    # --- SECTION 3 ---
    add_h1("3. SCHÉMA DES ENTITÉS ET BASE DE DONNÉES MVP")
    add_body("Pour soutenir l'ensemble des fonctionnalités du MVP, le modèle de données s'articule autour de 6 entités fondamentales :")

    # Table DB Entities
    tbl_db = doc.add_table(rows=7, cols=3)
    tbl_db.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_db.autofit = False
    
    db_data = [
        ("Nom de l'Entité", "Champs Clés Principaux", "Description & Rôle dans le MVP"),
        ("Users", "id, name, email, password_hash, role", "Stocke les comptes utilisateurs (Admin, Caissier, Agent, Client)."),
        ("Events", "id, title, date, location, image_url, status", "Gère la liste des matchs et événements sportifs programmés."),
        ("Venues / Seats", "id, venue_name, stand, sector, row, seat_num, status", "Représente les tribunes, secteurs et chaque siège numéroté du stade."),
        ("Prices", "id, event_id, stand_category, price_dhs", "Définit les tarifs associés aux tribunes pour chaque événement."),
        ("Tickets", "id, ticket_code, event_id, seat_id, buyer_name, qr_hash, status", "Billet unique généré avec QR Code Checksum et statut (VALIDE, SCANNE)."),
        ("AccessLogs", "id, ticket_id, scanned_at, agent_id, gate_name, result", "Historique de chaque scan d'accès avec horodatage et résultat (OK/FRAUDE).")
    ]

    col_w = [Inches(1.5), Inches(2.3), Inches(2.7)]
    for idx, r_data in enumerate(db_data):
        row = tbl_db.rows[idx]
        for c_i, val_txt in enumerate(r_data):
            cell = row.cells[c_i]
            cell.width = col_w[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val_txt)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(9.0)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)

    # --- SECTION 4 ---
    add_h1("4. PLAN D'EXÉCUTION & SPRINTS DE DÉVELOPPEMENT DU MVP")
    add_body("La concrétisation du MVP s'effectuera selon un plan de développement agile découpé en 4 Sprints successifs :")

    add_bullet("Mise en place du projet, structure des bases de données et système d'authentification RBAC.", "Sprint 1 — Core Infrastructure & Auth : ")
    add_bullet("Développement du visualiseur 2D de stade, sélection graphique des sièges et moteur d'événement/tarifs.", "Sprint 2 — Plan 2D & Moteur Billetterie : ")
    add_bullet("Développement du guichet POS tactile express (3 clics) et de la e-boutique web spectateur.", "Sprint 3 — Caisse POS & E-Commerce Web : ")
    add_bullet("Développement du scanner QR Code, mode Offline, dashboard de supervision et tests d'intégration.", "Sprint 4 — Scanner PDA, Mode Offline & Dashboard : ")

    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Fonctionnalites_MVP_E-Ticket_Pro.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_mvp_features_doc()
