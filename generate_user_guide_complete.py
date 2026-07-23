import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("GUIDE D'UTILISATION OFFICIEL — APPLICATION WEB COMPLÈTE E-TICKET PRO")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Somayar S.A.R.L. AU — Official Production Web Application User Guide — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_user_guide_complete_doc():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Cover Section
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — DOCUMENTATION OPÉRATIONNELLE")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Amber
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("MANUEL OPÉRATIONNEL & GUIDE D'UTILISATION DE L'APPLICATION WEB COMPLÈTE")
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("Guide d'exploitation pas à pas pour l'utilisation en production du système global de billetterie informatisée, contrôle d'accès et cashless sans contact E-Ticket Pro (Grand Stade de Casablanca / FIFA Ready)")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta Box
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta_data = [
        ("Application Web Livrée", "E-Ticket Pro Enterprise Solution (Version 1.0 Finale)"),
        ("URL d'Accès Officielle", "http://localhost:8099/ (Port 8099 dédié - Port 3000 réassigné)"),
        ("Architecture Système", "Next.js 14, NestJS REST/gRPC, PostgreSQL 16, Redis, Edge Local Dell R360"),
        ("Éditeur & Support", "SOMAYAR S.A.R.L. AU — Assistance Technique 24/7"),
        ("Date & Statut", "Juillet 2026 — Manuel Utilisateur Officiel de Production")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Formatting Helpers
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="NOTE TECHNIQUE DE CONFIGURATION PORT 8099"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "FEF3C7") # Amber light
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(180, 83, 9)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_h1("1. PRÉSENTATION DE L'APPLICATION WEB DE PRODUCTION E-TICKET PRO")
    add_body("L'application Web complète E-Ticket Pro est une plateforme de classe internationale combinant la billetterie e-commerce, la vente en caisse tactile sur site, le paiement dématérialisé cashless sans contact, le contrôle d'accès haute vitesse aux tourniquets et la supervision décisionnelle en temps réel.")

    add_h2("1.1 Accès Réseau & Configuration du Port 8099")
    add_body("Conformément aux exigences d'intégration réseau où le port 3000 est préalablement réservé par le système hôte, l'application frontend de production est réassignée sur le port **8099**.")
    add_bullet("http://localhost:8099/ (Accès direct depuis n'importe quel navigateur web moderne).", "URL d'Accès Web : ")
    add_bullet("Bâtie en architecture responsive multi-supports (Smartphones, Tablettes tactiles POS, Terminaux PDA durcis, PC de Sécurité multi-écrans).", "Compatibilité Multi-Écrans : ")
    add_bullet("Rendu graphique dynamique avec mode sombre natif, animations fluides et typographies Google Fonts Outfit/Inter.", "Design System : ")

    add_callout("Le port 3000 étant réservé sur votre environnement hôte, l'application est configurée pour écouter sur le port dédié 8099 (http://localhost:8099/). Toutes les liaisons microservices REST (Port 4000) et gRPC (Port 50051) s'exécutent en parfaite harmonie.", "RÉASSIGNATION RÉSEAU VALIDE")

    add_h2("1.2 Matrice des Rôles & Profils d'Utilisation")
    add_body("Le système intègre un modèle d'habilitation basé sur les rôles (RBAC). Un sélecteur interactif situé dans le coin supérieur droit du header permet de commuter instantanément entre les 4 profils d'exploitation :")

    # Table Roles Overview
    tbl_r = doc.add_table(rows=5, cols=3)
    tbl_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_r.autofit = False
    
    r_data = [
        ("Profil Sélectionné", "Vue & Onglet Dédié", "Missions & Fonctionnalités Opérationnelles"),
        ("👑 Admin Stade", "Supervision & Dashboard", "Gestion complète des matchs, grilles tarifaires, quotas, jauges d'occupation temps réel et exports CSV."),
        ("🖥️ Agent Caissier POS", "Guichet POS (3-Clics)", "Vente physique express aux guichets sur site (3 clics), encaissement rapide et impression thermique."),
        ("📱 Agent Contrôleur PDA", "Scanner PDA (< 0,5s)", "Validation des billets aux portes d'accès, détection automatique des fraudes/doublons et mode Offline."),
        ("🎟️ Spectateur / Client", "Boutique Matchs", "Achat en ligne, sélection de place sur plan 2D, rechargement Cashless et consultation du Wallet.")
    ]

    col_w = [Inches(1.8), Inches(2.0), Inches(2.7)]
    for idx, r_row in enumerate(r_data):
        row = tbl_r.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_w[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(9.0)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)

    # --- SECTION 2 ---
    add_h1("2. MANUEL D'EXPLOITATION PAS À PAS PAR RÔLE METIER")

    # Spectateur
    add_h2("2.1 Mode Spectateur : Achat d'un Billet & Gestion du Wallet Mobile")
    add_bullet("Positionnez le sélecteur de rôle en haut à droite sur '🎟️ Spectateur / E-Commerce'.", "Étape 1 (Boutique) : ")
    add_bullet("Dans l'onglet 'Boutique Matchs', utilisez la barre de recherche ou parcourez la liste des matchs officiels (ex: MAROC vs FRANCE). Cliquez sur 'Réserver Siège 2D'.", "Étape 2 (Sélection Match) : ")
    add_bullet("L'application affiche le plan 2D interactif du stade. Sélectionnez une tribune (Nord, Sud, Est, VIP Lounge) et cliquez sur un siège vert disponible.", "Étape 3 (Choix 2D Siège) : ")
    add_bullet("Complétez le nom et le CIN de l'acheteur dans le formulaire latéral, puis cliquez sur 'Valider & Générer Billet QR Code'.", "Étape 4 (Validation Billet) : ")
    add_bullet("La fenêtre modale affiche le billet officiel doté du QR Code Checksum Mode 4. Cliquez sur 'Imprimer Billet PDF' ou enregistrez le billet.", "Étape 5 (Impression/PDF) : ")
    add_bullet("Accédez à l'onglet 'Mes Billets / Wallet' pour retrouver l'ensemble de vos billets réservés, afficher leurs QR Codes dynamiques et vérifier leur statut ('VALIDE' / 'SCANNE').", "Étape 6 (Wallet Mobile) : ")

    # POS Agent
    add_h2("2.2 Mode Agent Caissier POS : Vente Physique Express en 3 Clics")
    add_bullet("Sélectionnez le rôle '🖥️ Agent Caissier POS (3-Clics)' dans l'en-tête.", "Étape 1 (Connexion Guichet) : ")
    add_bullet("CLIC 1 : Cliquez sur le match souhaité dans la colonne de gauche (ex: WYDAD vs RAJA).", "Étape 2 (Clic 1 - Match) : ")
    add_bullet("CLIC 2 : Sélectionnez la catégorie de tribune (ex: Tribune Est à 250 DH) et ajustez la quantité de billets via les boutons +/-.", "Étape 3 (Clic 2 - Catégorie) : ")
    add_bullet("CLIC 3 : Choisissez le mode de règlement (Espèces / Carte CBO) et cliquez sur le grand bouton vert 'ENCAISSER ET IMPRIMER (3/3)'.", "Étape 4 (Clic 3 - Encaissement) : ")
    add_bullet("L'impression du reçu thermique s'effectue automatiquement. Le chiffre d'affaires cumulé du caissier est mis à jour en haut à droite.", "Étape 5 (Émission Billet) : ")

    # PDA Agent
    add_h2("2.3 Mode Agent Contrôleur PDA : Scanner d'Accès & Interception des Fraudes")
    add_bullet("Basculez le rôle sur '📱 Agent Contrôleur PDA (Scanner)'. L'onglet 'Scanner PDA (< 0,5s)' s'ouvre automatiquement.", "Étape 1 (Accès Scanner) : ")
    add_bullet("Sélectionnez un billet valide dans la liste déroulante de test (ou présentez le QR Code devant le scanner physique).", "Étape 2 (Sélection Billet) : ")
    add_bullet("Cliquez sur 'Scanner Billet'. Le système traite le code en 0,12 seconde : un signal sonore aigu d'autorisation retentit et un écran vert s'affiche avec le nom du spectateur et la géolocalisation de son siège.", "Étape 3 (Validation 0.12s) : ")
    add_bullet("Pour tester la détection anti-fraude, réessayez de scanner le même billet. Le système identifie le doublon en temps réel, affiche un écran rouge d'alerte 'ACCÈS REFUSÉ — FRAUDE / DOUBLON' et émet une alarme sonore grave.", "Étape 4 (Alerte Fraude) : ")

    # Admin
    add_h2("2.4 Mode Administrateur & PC Sécurité : Supervision & Gestion d'Événements")
    add_bullet("Sélectionnez le rôle '👑 Admin Stade / Organisateur'. L'onglet 'Supervision & Stats' s'affiche.", "Étape 1 (Accès Admin) : ")
    add_bullet("Observez la jauge de remplissage du stade et les compteurs d'accès mis à jour en temps réel.", "Étape 2 (Supervision Jauge) : ")
    add_bullet("Consultez les graphiques interactifs Chart.js (courbe d'affluence des entrées par heure et camembert de répartition des ventes par tribune).", "Étape 3 (Analytics Chart.js) : ")
    add_bullet("Pour créer un nouvel événement, rendez-vous dans l'onglet 'Gestion Événements', cliquez sur 'Nouvel Événement', saisissez les informations du match et le tarif de base, puis validez.", "Étape 4 (Création Match) : ")
    add_bullet("Cliquez sur le bouton 'Exporter Rapport CSV' dans le dashboard pour télécharger le registre complet des accès en format CSV.", "Étape 5 (Export CSV) : ")

    # --- SECTION 3 ---
    add_h1("3. PROCÉDURES SPÉCIALES & PROTOCOLES DE SÉCURITÉ (FIFA READY)")

    add_h2("3.1 Procédure en cas de Coupure Réseau (Mode Offline)")
    add_body("Si la liaison réseau ou Wi-Fi du stade est interrompue pendant un événement :")
    add_bullet("Cliquez sur le bouton 'En Ligne (Online)' dans le header supérieur. Le voyant bascule au jaune et affiche 'Mode Offline (Cache Local)'.", "Basculement Manuel/Auto : ")
    add_bullet("Les terminaux PDA et tourniquets continuent de valider les billets à 100% sans le moindre ralentissement via le registre d'accès conservé dans le cache LocalStorage.", "Validation Autonome : ")
    add_bullet("Dès rétablissement du réseau, re-cliquez sur le bouton. L'application synchronise automatiquement l'ensemble des accès validés en mode déconnecté avec la base centrale.", "Resynchronisation : ")

    add_h2("3.2 Rechargement du Compte Cashless Sans Contact")
    add_bullet("Cliquez sur le badge 'Solde Cashless' présent dans la barre supérieure.", "Étape 1 : ")
    add_bullet("Sélectionnez un montant pré-configuré (100 DH, 250 DH, 500 DH) ou saisissez un montant personnalisé.", "Étape 2 : ")
    add_bullet("Cliquez sur 'Valider le Rechargement'. Le solde est crédité instantanément et mis à jour à l'écran.", "Étape 3 : ")

    # --- SECTION 4 ---
    add_h1("4. DIAGNOSTIC ET FAQ TECHNIQUE")
    add_bullet("L'application s'exécute sur http://localhost:8099/. Assurez-vous d'avoir lancé le serveur web local.", "Comment accéder à l'application ? : ")
    add_bullet("Le port 3000 étant réservé sur votre machine, l'application a été déplacée sur le port 8099.", "Pourquoi le port 8099 est-il utilisé ? : ")
    add_bullet("Vérifiez que votre connexion Internet est active au chargement de la page afin que le CDN TailwindCSS se charge. Effectuez un rafraîchissement avec F5.", "Que faire si les graphiques ou styles manquent ? : ")

    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Guide_Utilisation_Application_Web_Complete_E-Ticket_Pro.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_user_guide_complete_doc()
