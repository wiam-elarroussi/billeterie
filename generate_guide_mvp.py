import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("GUIDE D'UTILISATION & DOCUMENTATION MVP — E-TICKET PRO")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Somayar S.A.R.L. AU — E-Ticket Pro Express User Manual — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_guide_mvp_doc():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Cover Title
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — MANUEL OPÉRATIONNEL & TECHNIQUE")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Amber
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("GUIDE D'UTILISATION & DOCUMENTATION COMPLÈTE DU MVP")
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("Guide pas à pas des fonctionnalités opérationnelles de la solution Web MVP E-Ticket Pro (Billetterie 2D, Guichet POS 3-Clics, Scanner PDA Accès & Smart Dashboard)")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta Box
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta_data = [
        ("Application Produite", "E-Ticket Pro Express (MVP Web Application)"),
        ("URL d'Accès Local", "http://localhost:8099/ (Serveur Local Actif)"),
        ("Profils Couverts", "Super Admin, Agent POS Guichet, Agent PDA Scanner, Spectateur"),
        ("Conformité Système", "Standards FIFA Ready, Validation < 0,5s, Checksum Mode 4, Mode Offline"),
        ("Date & Version", "Juillet 2026 — Manuel d'Utilisation Officiel v1.0")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Formatting Helpers
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="CONSEIL D'UTILISATION RAPIDE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "F0FDF4") # Soft green
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(22, 101, 52)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_h1("1. PRÉSENTATION DE L'APPLICATION WEB MVP E-TICKET PRO")
    add_body("Le MVP (Minimum Viable Product) d'E-Ticket Pro est une application Web monopage (Single Page Application responsive) intégrant l'intégralité des briques fonctionnelles d'un système de billetterie et de contrôle d'accès de classe internationale.")
    
    add_h2("1.1 Architecture des Fichiers et Démarrage")
    add_bullet("Page HTML5 structurée intégrant le CDN TailwindCSS, les polices Google Fonts (Outfit & Inter), FontAwesome 6 et la bibliothèque graphique Chart.js.", "index.html : ")
    add_bullet("Feuille de style CSS personnalisée gérant les animations, les boutons interactifs, le plan de stade 2D et le format d'impression thermique.", "styles.css : ")
    add_bullet("Contrôleur JavaScript gérant l'état global de l'application, les 7 modules fonctionnels, le moteur de synthèse sonore (Web Audio API) et le stockage local.", "app.js : ")
    add_bullet("Serveur web HTTP local s'exécutant sur http://localhost:8099/.", "Serveur Local : ")

    add_h2("1.2 Prise en Main des 4 Profils Utilisateurs (RBAC)")
    add_body("Un sélecteur de rôle est situé dans le coin supérieur droit de l'en-tête. Il permet de tester instantanément l'application sous 4 perspectives différentes :")

    # Table Roles MVP
    tbl_r = doc.add_table(rows=5, cols=3)
    tbl_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_r.autofit = False
    
    r_data = [
        ("Profil Sélectionné", "Vue & Onglet par Défaut", "Droits & Actions Disponibles"),
        ("👑 Admin Stade", "Supervision & Dashboard", "Accès complet aux statistiques, gestion d'événements, export CSV et réinitialisation."),
        ("🖥️ Agent Caissier POS", "Guichet POS (3-Clics)", "Vente physique ultra-rapide aux guichets du stade, encaissement et impression immédiate."),
        ("📱 Agent Contrôleur PDA", "Scanner PDA (< 0,5s)", "Contrôle d'accès aux portes via scanner QR Code, simulation de fraudes et mode déconnecté."),
        ("🎟️ Spectateur / Client", "Boutique Matchs", "Achat e-commerce en ligne, sélection du siège sur plan 2D et consultation du wallet de billets.")
    ]

    col_w = [Inches(1.8), Inches(2.0), Inches(2.7)]
    for idx, r_row in enumerate(r_data):
        row = tbl_r.rows[idx]
        for c_i, val in enumerate(r_row):
            cell = row.cells[c_i]
            cell.width = col_w[c_i]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_i == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(9.0)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)

    # --- SECTION 2 ---
    add_h1("2. DESCRIPTION COMPLÈTE DES 7 MODULES INTEGRÉS")

    add_h2("2.1 Module 1 : Authentification & Sélecteur de Rôles RBAC")
    add_body("Ce module permet de simuler la connexion des différents acteurs. En modifiant la valeur du sélecteur en haut à droite, l'application bascule automatiquement sur l'onglet pertinent et adapte l'avatar et le nom de l'utilisateur actif.")

    add_h2("2.2 Module 2 : Plan 2D de Stade & Réservation Interactive des Sièges")
    add_body("Ce module offre un rendu schématique 2D d'une enceinte sportive avec sa pelouse centrale et ses 4 tribunes principales :")
    add_bullet("Tribune Nord (Catégorie 2 — 150 DH)", "• Tribune Nord : ")
    add_bullet("Tribune Sud (Catégorie 2 — 150 DH)", "• Tribune Sud : ")
    add_bullet("Tribune Est (Catégorie 1 — 250 DH)", "• Tribune Est : ")
    add_bullet("VIP Lounge (Prestige — 600 DH)", "• VIP Lounge : ")
    add_body("Pour chaque tribune, une grille de 24 sièges est générée interactivement. Le code couleur permet d'identifier l'état de chaque siège : Vert (Libre), Rouge (Occupé), Bleu (Sélectionné) et Or (VIP). Le panneau de droite affiche instantanément le rang, le numéro de siège, la tribune et le tarif unitaire.")

    add_h2("2.3 Module 3 : Gestion d'Événements & QR Code Checksum Mode 4")
    add_body("Ce module gère le catalogue des matchs et la génération des titres d'accès. Chaque billet généré possède un identifiant unique incrackable (ex: ETK-2026-X9F4A7B2) intégrant un caractère de Checksum en Mode 4. Il peut être téléchargé ou imprimé au format thermique/PDF.")

    add_h2("2.4 Module 4 : Guichet de Vente POS Express (Mode 3-Clics)")
    add_body("Le guichet POS est optimisé pour les écrans tactiles afin de réduire au strict minimum le temps d'encaissement en période d'affluence :")
    add_bullet("Sélection du match dans la liste.", "Clic 1 : ")
    add_bullet("Sélection de la tribune (Nord, Sud, Est, VIP) et de la quantité.", "Clic 2 : ")
    add_bullet("Encaissement (Espèces/CB) et génération instantanée du billet.", "Clic 3 : ")

    add_h2("2.5 Module 5 : Boutique E-Commerce Spectateur & Wallet Billet")
    add_body("Ce module reproduit l'expérience d'achat d'un supporter sur smartphone. Il inclut la présentation des affiches de match, la sélection du siège en ligne et la section 'Mes Billets' qui affiche le QR Code dynamique et le statut de chaque réservation.")

    add_h2("2.6 Module 6 : Scanner PDA & Contrôle d'Accès (< 0,5s & Mode Offline)")
    add_body("Le module scanner simule la validation des billets aux tourniquets du stade. Il se caractérise par :")
    add_bullet("Temps d'exécution moyen de 0,12 seconde par billet.", "Validation Ultra-Rapide : ")
    add_bullet("Écran vert avec bip aigu pour accès accordé vs Écran rouge avec alarme grave pour billet déjà scanné (doublon) ou invalide.", "Feedback Visuel & Sonore : ")
    add_bullet("Un bouton 'En Ligne / Mode Offline' dans l'en-tête permet de basculer la validation sur le cache local LocalStorage pour simuler une coupure de réseau Wi-Fi.", "Gestion Déconnectée (Offline) : ")

    add_h2("2.7 Module 7 : Supervision & Smart Dashboard PC Sécurité")
    add_body("Le dashboard de supervision centralise l'état opérationnel du stade en temps réel : jauge de remplissage circulaire, chiffre d'affaires cumulé, nombre de spectateurs entrés, nombre de fraudes interceptées et graphiques interactifs Chart.js (flux des entrées par heure et répartition des ventes par tribune).")

    add_callout("Pour tester le mode déconnecté (Offline), cliquez sur le bouton 'En Ligne (Online)' dans l'en-tête supérieur. Le système bascule en mode 'Offline (Cache Local)' et continue de valider les billets sans la moindre interruption !", "TEST DU MODE OFFLINE DANS LE MVP")

    # --- SECTION 3 ---
    add_h1("3. GUIDE PAS À PAS D'UTILISATION (SCÉNARIOS PAR PROFIL)")

    add_h2("3.1 Scénario 1 : Achat d'un Billet par un Spectateur (Client)")
    add_bullet("Vérifiez que le sélecteur de rôle en haut à droite est positionné sur '🎟️ Spectateur / E-Commerce'.", "Étape 1 : ")
    add_bullet("Dans l'onglet 'Boutique Matchs', cliquez sur le bouton 'Réserver Siège 2D' du match MAROC vs FRANCE.", "Étape 2 : ")
    add_bullet("Sur le plan 2D du stade, choisissez une tribune (ex: Tribune Nord) et cliquez sur un siège vert disponible.", "Étape 3 : ")
    add_bullet("Renseignez le nom de l'acheteur dans le formulaire latéral (ex: Karim Bennani).", "Étape 4 : ")
    add_bullet("Cliquez sur 'Valider & Générer Billet QR Code'. Le billet officiel s'affiche avec son QR Code scannable.", "Étape 5 : ")
    add_bullet("Rendez-vous dans l'onglet 'Mes Billets / Wallet' pour consulter le billet enregistré et son statut 'VALIDE'.", "Étape 6 : ")

    add_h2("3.2 Scénario 2 : Vente Express au Guichet POS (Caissier)")
    add_bullet("Sélectionnez le rôle '🖥️ Agent Caissier POS (3-Clics)' ou ouvrez l'onglet 'Guichet POS (3-Clics)'.", "Étape 1 : ")
    add_bullet("Clic 1 : Choisissez le match 'WYDAD CASABLANCA vs RAJA CA' dans la colonne de gauche.", "Étape 2 : ")
    add_bullet("Clic 2 : Cliquez sur la catégorie 'Tribune Est (250 DH)' et ajustez la quantité à 2 billets.", "Étape 3 : ")
    add_bullet("Clic 3 : Cliquez sur le grand bouton vert 'ENCAISSER ET IMPRIMER (3/3)'.", "Étape 4 : ")
    add_bullet("La fenêtre modale du billet thermique s'affiche instantanément. Cliquez sur 'Imprimer Billet PDF' pour lancer l'impression.", "Étape 5 : ")

    add_h2("3.3 Scénario 3 : Contrôle d'Accès aux Portes & Interception de Fraude (Agent PDA)")
    add_bullet("Positionnez le rôle sur '📱 Agent Contrôleur PDA (Scanner)' ou ouvrez l'onglet 'Scanner PDA (< 0,5s)'.", "Étape 1 : ")
    add_bullet("Dans la liste déroulante sous le scanner, sélectionnez un billet dont le statut est 'VALIDE'.", "Étape 2 : ")
    add_bullet("Cliquez sur 'Scanner Billet'. Le scanner valide le billet en 0.12s, affiche un écran vert et émet un bip sonore d'autorisation.", "Étape 3 : ")
    add_bullet("Cliquez une seconde fois sur 'Scanner Billet' pour simuler la réutilisation frauduleuse du même billet par une autre personne.", "Étape 4 : ")
    add_bullet("Le système détecte immédiatement le doublon, affiche un écran rouge d'alerte 'ACCÈS REFUSÉ — FRAUDE / DOUBLON' et émet une alarme sonore.", "Étape 5 : ")

    add_h2("3.4 Scénario 4 : Supervision du Stade & Export de Rapport (Admin / PC Sécurité)")
    add_bullet("Sélectionnez le rôle '👑 Admin Stade / Organisateur' ou ouvrez l'onglet 'Supervision & Stats'.", "Étape 1 : ")
    add_bullet("Observez la jauge de remplissage du stade et les compteurs d'entrées mis à jour en temps réel.", "Étape 2 : ")
    add_bullet("Consultez le graphique des flux d'entrées par heure et le diagramme circulaire des ventes par tribune.", "Étape 3 : ")
    add_bullet("Cliquez sur le bouton 'Exporter Rapport CSV' en haut à droite pour télécharger la synthèse des accès.", "Étape 4 : ")

    # --- SECTION 4 ---
    add_h1("4. FAQ & DÉPANNAGE TECHNIQUE")
    add_bullet("Ouvrez simplement le lien http://localhost:8099/ dans votre navigateur Chrome, Edge ou Firefox. Assurez-vous que le serveur local est actif.", "Comment ouvrir l'application MVP ? : ")
    add_bullet("Si les styles apparaissent déformés, vérifiez votre connexion Internet afin que le CDN TailwindCSS se charge correctement, puis faites un rafraîchissement avec F5.", "Que faire si les styles CSS ne s'affichent pas ? : ")
    add_bullet("Dans l'onglet Scanner PDA, sélectionnez le même billet deux fois de suite. La seconde tentative déclenchera la détection de doublon.", "Comment tester la sécurité anti-fraude ? : ")

    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Guide_Utilisation_et_Documentation_MVP.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_guide_mvp_doc()
