import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color of a cell (hex string without #)."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("DÉFINITION DES BESOINS & UTILISATEURS CIBLES — E-TICKET PRO")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("SOMAYAR S.A.R.L. AU — Solution E-Ticket Pro / Standards FIFA Ready — Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(120, 120, 120)

def create_besoins_utilisateurs_doc():
    doc = docx.Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Title Page / Cover
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(30)
    
    p_org = doc.add_paragraph()
    r_org = p_org.add_run("SOMAYAR S.A.R.L. AU — ÉQUIPEMENTS DE CONTRÔLE & DE SÉCURITÉ")
    r_org.font.bold = True
    r_org.font.size = Pt(11)
    r_org.font.color.rgb = RGBColor(217, 119, 6) # Gold / Amber
    
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(10)
    r_title = p_title.add_run("DOCUMENT DE DÉFINITION DES BESOINS & CARTOGRAPHIE DES UTILISATEURS CIBLES")
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("Analyse détaillée des exigences fonctionnelles, opérationnelles et des profils utilisateurs de la solution globale de billetterie, contrôle d'accès et cashless E-Ticket Pro (Standards FIFA Ready)")
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    # Meta table
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False
    
    meta = [
        ("Nom du Projet", "Écosystème E-Ticket Pro (Billetterie, Cashless, Accès, Report)"),
        ("Contexte du Marché", "Appel d'Offres N° 02/2026/GSC — Grand Stade de Casablanca"),
        ("Document Produit", "Définition des Besoins Métiers & Cartographie Utilisateurs Cibles"),
        ("Éditeur / Integrateur", "SOMAYAR S.A.R.L. AU"),
        ("Date & Statut", "Juillet 2026 — Version 1.0 Définitive")
    ]
    
    for idx, (label, val) in enumerate(meta):
        row = table_meta.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.3)
        c1.width = Inches(4.2)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(27, 54, 93)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        
        set_cell_background(c0, "F1F5F9")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    doc.add_page_break()
    
    # Helper formatting functions
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(27, 54, 93)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(217, 119, 6)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(74, 85, 104)
        return h

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        r = p.add_run(text)
        r.font.italic = italic
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(27, 54, 93)
        p.add_run(text)
        return p

    def add_callout(text, title="EXIGENCE DE CONCEPTION & BESOIN"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.5)
        set_cell_background(c, "EFF6FF") # Soft Blue
        set_cell_margins(c, top=120, bottom=120, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(title + "\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(30, 64, 175)
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = RGBColor(50, 50, 50)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- SECTION 1 ---
    add_h1("1. INTRODUCTION & CONTEXTE DE LA DÉFINITION DES BESOINS")
    add_body("Le présent document a pour objet de formaliser de manière exhaustive l'ensemble des besoins fonctionnels, techniques et opérationnels liés à la mise en œuvre de la solution globale de billetterie, de paiement cashless et de contrôle d'accès E-Ticket Pro. Il établit également la cartographie précise des profils et utilisateurs cibles appelés à interagir avec le système.")
    
    add_h2("1.1 Contexte et Enjeux Métiers")
    add_body("La gestion de grands événements sportifs et culturels au sein d'infrastructures d'envergure internationale (telles que le Grand Stade de Casablanca) impose de traiter des volumes massifs de spectateurs dans des conditions de sécurité, de rapidité et de confort optimales.")
    add_bullet("Permettre un écoulement des foules sans congestion aux portes d'accès (impératif de contrôle en moins de 0,5 seconde).", "Fluidité Massique : ")
    add_bullet("Éliminer la falsification des billets, la revente illicite et l'utilisation de doublons.", "Sécurisation Anti-Fraude : ")
    add_bullet("Diversifier les canaux de vente (en ligne, guichets sur site, terminaux tiers) tout en conservant un contrôle centralisé des jauges.", "Vente Omnicanale : ")
    add_bullet("Supprimer la manipulation d'espèces sur site grâce à une solution dématérialisée Cashless sécurisée.", "Expérience Cashless : ")
    add_bullet("Garantir la continuité des opérations même en cas d'interruption du réseau informatique ou d'alimentation secteur.", "Résilience Télécom & Autonomie Offline : ")

    add_h2("1.2 Analyse des Problématiques Actuelles (Pain Points)")
    add_body("L'analyse des systèmes traditionnels ou non intégrés met en évidence plusieurs axes critiques de défaillance auxquels la solution E-Ticket Pro apporte une réponse ciblée :")

    # Table Pain Points vs Solutions
    tbl_pain = doc.add_table(rows=6, cols=2)
    tbl_pain.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pain.autofit = False
    
    pain_data = [
        ("Point de Douleur / Problématique Constatée", "Réponse Métier & Solution E-Ticket Pro"),
        ("Files d'attente interminables aux tourniquets dues à des temps de lecture lents.", "Validation ultra-rapide (< 0,5s par billet) multi-technologies (NFC, QR, Barcode 1D/2D)."),
        ("Vente de faux billets / duplicatas imprimés frauduleusement.", "Code sécurisé unique avec Checksum Mode 4 infalsifiable et invalidation instantanée après 1er scan."),
        ("Paralysie des accès en cas de coupure de réseau ou de serveur distant.", "Mode Offline résilient avec stockage local synchronisé et validation en autonomie totale."),
        ("Lenteur des ventes en guichet aux heures de pointe avant le coup d'envoi.", "Interface POS tactile ultra-ergonomique permettant la vente physique en 3 clics seulement."),
        ("Opacité sur l'occupation réelle du stade et manque de données en temps réel.", "Tableaux de bord dynamiques et monitoring graphique du PC Sécurité accessibles 24/7 sur PC & Mobile.")
    ]

    for idx, (p_txt, s_txt) in enumerate(pain_data):
        row = tbl_pain.rows[idx]
        ca, cb = row.cells[0], row.cells[1]
        ca.width = Inches(2.5)
        cb.width = Inches(4.0)
        pa, pb = ca.paragraphs[0], cb.paragraphs[0]
        ra, rb = pa.add_run(p_txt), pb.add_run(s_txt)
        if idx == 0:
            ra.font.bold = True
            ra.font.color.rgb = RGBColor(255, 255, 255)
            rb.font.bold = True
            rb.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(ca, "1B365D")
            set_cell_background(cb, "1B365D")
        else:
            ra.font.bold = True
            ra.font.size = Pt(9.5)
            ra.font.color.rgb = RGBColor(27, 54, 93)
            rb.font.size = Pt(9.5)
            bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
            set_cell_background(ca, bg)
            set_cell_background(cb, bg)
        set_cell_margins(ca, top=80, bottom=80, left=100, right=100)
        set_cell_margins(cb, top=80, bottom=80, left=100, right=100)

    # --- SECTION 2 ---
    add_h1("2. DÉFINITION DÉTAILLÉE DES BESOINS SYSTEME")

    add_h2("2.1 Besoins Opérationnels de Contrôle d'Accès")
    add_bullet("Le système doit valider tout titre d'accès présent sur support papier ou écran de smartphone en moins de 0,5 seconde.", "Vitesse d'Exécution : ")
    add_bullet("Capacité à lire les codes-barres 1D (EAN/Code128), 2D (QR Code, DataMatrix), puces sans contact NFC, RFID ainsi que les portefeuilles Apple Pay et Google Pay.", "Multi-Support : ")
    add_bullet("En cas de rupture de communication réseau, les obstacles d'accès physiques (tourniquets) et les PDA portables doivent poursuivre la validation locale sans interruption.", "Autonomie Déconnectée (Offline) : ")
    add_bullet("Capacité d'interdire l'accès aux personnes signalées via l'importation dynamique de listes noires (et contrôle via listes blanches).", "Gestion des Listes Noires/Blanches : ")

    add_h2("2.2 Besoins de Gestion de Billetterie & Paramétrage (E-Ticket Pro)")
    add_bullet("Outil graphique 2D de dessin de l'enceinte permettant de définir graphiquement les entrées, portes, tribunes, zones, secteurs, rangs et sièges numérotés.", "Éditeur de Plan d'Enceinte 2D : ")
    add_bullet("Définition du sens de passage recommandé de siège à siège lors de l'attribution des places.", "Orientation & Sens de Passage : ")
    add_bullet("Possibilité de définir des tarifs par catégorie de place, type de visiteur, phase de compétition (éliminatoires, finales) et formules abonnement.", "Grille Tarifaire Dynamique : ")
    add_bullet("Définition stricte du nombre maximal de billets par impression (ex: 3 billets max par acheteur) et attribution de quotas fermes par guichet ou partenaire.", "Gestion Stricte des Quotas : ")
    add_bullet("Éditeur graphique permettant de personnaliser la charte visuelle des billets imprimés (thermiques), des e-tickets PDF et des badges d'accréditation.", "Personnalisation des Billets & Badges : ")

    add_h2("2.3 Besoins de Paiement Cashless & E-Commerce (E-Ticket Pay)")
    add_bullet("Intégration d'un module de billetterie en ligne autonome ou intégrable dans un site web existant via API.", "Boutique Web & E-Commerce : ")
    add_bullet("Guichet de caisse tactile permettant à l'opérateur de réaliser l'ensemble d'une transaction de vente physique en 3 clics.", "Guichet POS Tactile Express : ")
    add_bullet("Système dématérialisé Cashless permettant d'associer un solde monétaire au billet ou badge du spectateur, rechargeable en ligne ou sur site.", "Module Cashless Dématérialisé : ")
    add_bullet("Module permettant au spectateur de visualiser et de retrouver son siège sur la carte 2D interactive depuis le navigateur Web de son smartphone.", "Repérage 2D du Siège sur Mobile : ")

    add_h2("2.4 Besoins d'Interopérabilité & Normes FIFA Ready")
    add_bullet("Disponibilité de SDK et d'APIs RESTful sécurisées pour l'interfaçage natif avec les systèmes centraux de billetterie FIFA Ticketing ou Ticketmaster.", "APIs & SDK FIFA Ready : ")
    add_bullet("Gestion fluide des accréditations pour les populations spécifiques (VIP, VVIP, Presse, Bénévoles, Délégations officiels).", "Accréditations Spécifiques : ")
    add_bullet("Implémentation de caractères de Checksum en Mode 4 rendant tout code-barres généré incrackable et infalsifiable.", "Anti-Fraude Checksum Mode 4 : ")

    add_h2("2.5 Besoins Décisionnels & Smart Reporting (E-Ticket Report)")
    add_bullet("Affichage en temps réel des statistiques de vente, des revenus et des taux d'occupation des jauges par porte et par tribune.", "Tableaux de Bord Dynamiques : ")
    add_bullet("Monitoring graphique centralisé depuis le PC de Sécurité permettant de visualiser l'état de fonctionnement et le flux de chaque tourniquet.", "Monitoring PC Sécurité : ")
    add_bullet("Rapports détaillés sur le profil des acheteurs (origine géographique, âge, genre, canal d'acquisition marketing).", "Analytics & Démographie : ")
    add_bullet("Génération et export en 1 clic de rapports sous formats Excel, Word, CSV, PDF et XML.", "Exports Multi-Formats : ")
    add_bullet("Application mobile de supervision en temps réel compatible Android, iOS et macOS.", "Supervision Mobile : ")

    add_callout("Le serveur central s'appuie sur une architecture matérielle Dell PowerEdge R360 (Intel Xeon, 128 Go RAM DDR5, BOSS-N1 HWRAID SSD NVMe, double alimentation Hot-Swap 700W Titanium), garantissant un taux de disponibilité minimal de 99,99% pendant la tenue des événements.", "BESOIN CRITIQUE D'INFRASTRUCTURE RACK 1U")

    # --- SECTION 3 ---
    add_h1("3. CARTOGRAPHIE DES UTILISATEURS CIBLES & PERSONAS")
    add_body("Pour garantir une adéquation parfaite entre l'ergonomie applicative et les besoins réels du terrain, la solution E-Ticket Pro identifie 6 profils d'utilisateurs cibles majeurs.")

    # Persona 1
    add_h2("3.1 Persona 1 : Karine M. — Superviseur PC Sécurité & Contrôle d'Accès")
    add_bullet("Superviser en temps réel l'ensemble des accès physiques du stade, détecter les tentatives de fraude ou les engorgements aux portes, et réagir immédiatement.", "Rôle & Objectifs : ")
    add_bullet("5 ans d'expérience en sécurité événementielle. Utilise le PC de sécurité avec 4 écrans de contrôle.", "Profil & Contexte : ")
    add_bullet("Interface graphique du PC de Sécurité E-Ticket Report, alertes en temps réel, monitoring des tourniquets.", "Outils Utilisés : ")
    add_bullet("Vue synthétique en un coup d'œil de la jauge globale, alerte visuelle si un tourniquet rencontre un blocage, injection rapide d'une liste noire d'urgence.", "Besoins Clés : ")

    # Persona 2
    add_h2("3.2 Persona 2 : Youssef K. — Responsable Billetterie & Chef de Projet Événements")
    add_bullet("Configurer les événements, programmer la grille tarifaire, affecter les quotas de billets par canal de vente et suivre le chiffre d'affaires.", "Rôle & Objectifs : ")
    add_bullet("Cadre commercial & opérationnel, gère la programmation de la saison sportive et des concerts.", "Profil & Contexte : ")
    add_bullet("Backoffice E-Ticket Pro (PC portable / poste de bureau), module de dessin 2D de stade.", "Outils Utilisés : ")
    add_bullet("Flexibilité dans la modification des tarifs, gestion simple des abonnés, édition rapide de rapports financiers pour la direction.", "Besoins Clés : ")

    # Persona 3
    add_h2("3.3 Persona 3 : Mehdi B. — Agent de Caisse POS (Guichetier sur Site)")
    add_bullet("Vendre des billets physiques au guichet du stade le jour du match, encaisser les paiements (Espèces/CB/Cashless) et réimprimer des commandes si besoin.", "Rôle & Objectifs : ")
    add_bullet("Opérateur de caisse temporaire ou permanent, travaillant dans un environnement bruyant à fort stress.", "Profil & Contexte : ")
    add_bullet("Écran tactile POS + Imprimante thermique de billets + Lecteur/Douchette de code-barres.", "Outils Utilisés : ")
    add_bullet("Interface tactile ultra-simplifiée (vente en 3 clics max), encaissement rapide, zéro bug lors des pics d'affluence avant le coup d'envoi.", "Besoins Clés : ")

    # Persona 4
    add_h2("3.4 Persona 4 : Samira L. — Agent de Contrôle aux Portes (Opératrice PDA)")
    add_bullet("Accueillir les spectateurs aux tourniquets et portillons, valider les billets nominatifs et les accréditations VIP/Presse via terminal mobile.", "Rôle & Objectifs : ")
    add_bullet("Agent de sécurité / contrôle mobile en posture debout aux entrées du stade.", "Profil & Contexte : ")
    add_bullet("Terminal mobile PDA durci avec scanner laser et lecteur NFC.", "Outils Utilisés : ")
    add_bullet("Validation instantanée (signal vert/rouge clair avec retour sonore), basculement automatique en mode Offline si le Wi-Fi faiblit.", "Besoins Clés : ")

    # Persona 5
    add_h2("3.5 Persona 5 : Thomas V. — Spectateur / Supporter (Utilisateur Final)")
    add_bullet("Acheter un billet en ligne sans contrainte, accéder rapidement au stade sans faire la queue, utiliser le solde Cashless et trouver son siège facilement.", "Rôle & Objectifs : ")
    add_bullet("Grand public, adepte de technologie mobile, exigeant sur l'expérience globale et la rapidité.", "Profil & Contexte : ")
    add_bullet("Smartphone (Navigateur Web mobile, E-Ticket PDF, Apple Wallet / Google Pay).", "Outils Utilisés : ")
    add_bullet("Achat fluide sur smartphone, accès sans contact, repérage 2D de son siège sur la carte du stade depuis son mobile.", "Besoins Clés : ")

    # Persona 6
    add_h2("3.6 Persona 6 : M. El Alami — Directeur des Opérations & Analyste Stratégique")
    add_bullet("Analyser les performances financières et opérationnelles des événements, évaluer les habitudes d'achat et valider la conformité aux exigences FIFA.", "Rôle & Objectifs : ")
    add_bullet("Membre de la direction, décideur stratégique et garant des engagements de l'infrastructure du stade.", "Profil & Contexte : ")
    add_bullet("Application mobile E-Ticket Report sur tablette/smartphone (iOS/Android/macOS) et synthèses PDF/Excel.", "Outils Utilisés : ")
    add_bullet("Données consolidées fiables, indicateurs de performance (KPIs) en temps réel, exports faciles pour le conseil d'administration.", "Besoins Clés : ")

    # Table Summary of User Roles
    add_h3("Synthèse Comparative des Profils Utilisateurs")
    tbl_users = doc.add_table(rows=7, cols=4)
    tbl_users.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_users.autofit = False
    
    users_data = [
        ("Profil Utilisateur", "Canal / Interface Principal", "Niveau de Fréquence", "Exigence Majeure"),
        ("Superviseur PC Sécurité", "Poste PC Sécurité / E-Ticket Report", "Temps réel continu", "Monitoring graphique & réactivité < 1s"),
        ("Responsable Billetterie", "Backoffice E-Ticket Pro (PC)", "Quotidienne / Hebdo", "Gestion des tarifs, quotas & plans 2D"),
        ("Agent de Caisse POS", "Guichet Tactile POS + Imprimante", "Jours d'événement (Intensif)", "Vente express en 3 clics tactiles"),
        ("Agent de Contrôle PDA", "Terminal Mobile PDA (NFC/Laser)", "Jours d'événement (Intensif)", "Validation < 0.5s & Mode Offline"),
        ("Spectateur / Supporter", "Boutique Web Mobile / Smartphone", "Occasionnelle", "Achat rapide & Repérage 2D du siège"),
        ("Directeur des Opérations", "App Mobile E-Ticket Report (iOS/Android)", "Périodique & Décisionnelle", "Tableaux de bord & KPIs consolidés")
    ]

    col_widths = [Inches(1.8), Inches(2.2), Inches(1.2), Inches(1.8)]
    for idx, row_data in enumerate(users_data):
        row = tbl_users.rows[idx]
        for c_idx, cell_txt in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(cell_txt)
            if idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1B365D")
            else:
                if c_idx == 0:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(27, 54, 93)
                r.font.size = Pt(9.0)
                bg = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=80, right=80)

    # --- SECTION 4 ---
    add_h1("4. CAS D'USAGE CRITIQUES & PARCOURS UTILISATEURS")
    
    add_h2("4.1 Parcours 1 : Achat d'un Ticket en Ligne et Accès au Stade par le Spectateur")
    add_bullet("Le spectateur accède à la e-boutique E-Ticket Pay depuis son smartphone.", "Étape 1 (Achat) : ")
    add_bullet("Il choisit son événement et sélectionne son siège sur la carte 2D dynamique de la tribune.", "Étape 2 (Sélection) : ")
    add_bullet("Il règle sa commande en ligne et reçoit instantanément son e-ticket doté du QR Code sécurisé Checksum Mode 4.", "Étape 3 (Paiement) : ")
    add_bullet("Le jour du match, à l'approche de la porte, il ouvre la carte 2D mobile pour repérer exactement son siège.", "Étape 4 (Repérage) : ")
    add_bullet("Présentation du smartphone sur le lecteur du tourniquet : validation et déverrouillage en moins de 0,5 seconde.", "Étape 5 (Entrée) : ")

    add_h2("4.2 Parcours 2 : Gestion d'une Coupure Réseau par l'Agent de Contrôle PDA")
    add_bullet("Le réseau Wi-Fi de la porte Nord subit une défaillance technique imprévue.", "Étape 1 (Incident) : ")
    add_bullet("Les terminaux PDA et l'unité de contrôle du tourniquet détectent la perte de liaison et basculent automatiquement en mode Offline.", "Étape 2 (Basculement) : ")
    add_bullet("L'agent PDA continue de scanner les billets sans moindre ralentissement grâce à la base locale synchronisée.", "Étape 3 (Continuité) : ")
    add_bullet("Dès le rétablissement de la connexion, les données de comptage et de validation sont resynchronisées en tâche de fond avec le serveur Dell R360.", "Étape 4 (Resynchronisation) : ")

    # --- SECTION 5 ---
    add_h1("5. MATRICE D'APPROBATION DE LA DÉFINITION DES BESOINS")
    add_body("La présente définition des besoins a été établie conjointement par l'équipe d'ingénierie système et les représentants métier de SOMAYAR S.A.R.L. AU et constitue la référence fonctionnelle du projet E-Ticket Pro.")

    tbl_appr = doc.add_table(rows=2, cols=2)
    tbl_appr.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_appr.autofit = False
    
    c_a0 = tbl_appr.rows[0].cells[0]
    c_a1 = tbl_appr.rows[0].cells[1]
    c_a0.width = Inches(3.25)
    c_a1.width = Inches(3.25)
    
    p_a0 = c_a0.paragraphs[0]
    p_a0.add_run("Pour la Direction des Opérations / Client :\n\n___________________________________\nNom, Titre & Signature").font.size = Pt(9.5)
    
    p_a1 = c_a1.paragraphs[0]
    p_a1.add_run("Pour l'Équipe Projet SOMAYAR S.A.R.L. AU :\n\n___________________________________\nChef de Projet E-Ticket Pro").font.size = Pt(9.5)
    
    set_cell_background(c_a0, "F8FAFC")
    set_cell_background(c_a1, "F8FAFC")
    set_cell_margins(c_a0, top=120, bottom=120, left=120, right=120)
    set_cell_margins(c_a1, top=120, bottom=120, left=120, right=120)

    add_header_footer(doc)
    
    output_path = "D:/New folder/smr/concp/1268 solution billeterie/Definition_Besoins_Utilisateurs_Cibles.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_besoins_utilisateurs_doc()
